"""创建带结构化元数据的Word知识库"""
from docx import Document

doc = Document()
doc.add_heading('贝壳公司员工手册', 0)
doc.add_paragraph('版本: 2024年第1版')
doc.add_paragraph('发布日期: 2024-01-01')
doc.add_paragraph('---')

# 第一章
doc.add_heading('第一章 差旅管理制度', 1)
doc.add_heading('1.1 出差补贴标准', 2)
doc.add_paragraph('【规定】贝壳员工出差补贴标准为500元/人/天,包含住宿和餐饮费用。')
doc.add_paragraph('【适用范围】全体正式员工')
doc.add_paragraph('【生效日期】2024年1月1日')
doc.add_paragraph('【文件编号】BK-HR-2024-001')

# 第二章
doc.add_heading('第二章 薪酬福利制度', 1)
doc.add_heading('2.1 应届生住房补贴', 2)
doc.add_paragraph('【规定】应届生入职贝壳时,公司提供2500元/月的住房补贴,随工资发放,连续12个月,总计30000元。')
doc.add_paragraph('【申请条件】应届毕业生(毕业2年内)')
doc.add_paragraph('【生效日期】2024年1月1日')
doc.add_paragraph('【文件编号】BK-HR-2024-002')

doc.add_heading('2.2 五险一金缴纳', 2)
doc.add_paragraph('【规定】公司按照北京市标准为员工缴纳五险一金,公积金缴纳比例为12%。')
doc.add_paragraph('【文件编号】BK-HR-2024-003')

# 第三章
doc.add_heading('第三章 假期管理制度', 1)
doc.add_heading('3.1 年假标准', 2)
doc.add_paragraph('【规定】年假标准如下:')
doc.add_paragraph('- 工作1-5年: 5天年假')
doc.add_paragraph('- 工作5-10年: 10天年假')
doc.add_paragraph('- 工作10年以上: 15天年假')
doc.add_paragraph('【文件编号】BK-HR-2024-004')

doc.add_heading('3.2 病假制度', 2)
doc.add_paragraph('【规定】员工因病需要休假的,凭医院证明可申请病假,病假期间工资按80%发放。')
doc.add_paragraph('【文件编号】BK-HR-2024-005')

doc.save('knowledge.docx')
print('✅ 结构化Word文档创建成功: knowledge.docx')
print('📋 包含文件编号、章节结构、生效日期等元数据')

